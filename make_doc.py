from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2B579A', qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 데이터
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '맑은 고딕'
            if ri % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F2F2', qn('w:val'): 'clear'
                })
                shading.append(shading_elm)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10)
        p.add_run(text).font.name = '맑은 고딕'
        p.runs[-1].font.size = Pt(10)
    else:
        p.text = text
        for run in p.runs:
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10)
    return p

# ================================================================
# 표지
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(100)
run = p.add_run('KI-RnB Pakistan (HNMPL)\n')
run.font.size = Pt(28)
run.font.bold = True
run.font.name = '맑은 고딕'
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

run2 = p.add_run('LX3 차종 추가 및 에러 방어 코드 작업 보고서\n\n')
run2.font.size = Pt(16)
run2.font.name = '맑은 고딕'
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

run3 = p.add_run('2026-08-15\n')
run3.font.size = Pt(14)
run3.font.name = '맑은 고딕'
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

run4 = p.add_run('프로젝트: E:\\파키스탄 LX3\\KI-RnB_1208\nGitHub: rr8602/KI-RnB-PK (private)')
run4.font.size = Pt(10)
run4.font.name = '맑은 고딕'
run4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ================================================================
# 1. 작업 개요
# ================================================================
add_heading('1. 작업 개요', 1)

doc.add_paragraph(
    '파키스탄 HNMPL 공장의 KI-RnB ABS/ESC 롤러벤치 검사 프로그램에 대해 '
    'LX3 HEV/ICE 차종 추가 및 런타임 에러 방어 코드를 전면 적용하였다.'
)

add_table(
    ['항목', '내용'],
    [
        ['대상 프로젝트', 'KI-RnB_1208 (파키스탄 LX3)'],
        ['작업 기간', '2026-08-15'],
        ['커밋 내역', '3건 (Initial → 에러방어 1차 → 에러방어 2차)'],
        ['수정 파일', '25개 .cs 파일 + 설정 파일'],
        ['변경 라인', '+800 / -71'],
        ['백업 브랜치', 'backup/before-error-fix'],
    ],
    [4, 12]
)

# ================================================================
# 2. 커밋 히스토리
# ================================================================
doc.add_page_break()
add_heading('2. 커밋 히스토리', 1)

add_table(
    ['#', '시간', '커밋 메시지', '주요 내용'],
    [
        ['1', '13:37', 'Initial commit: KI-RnB Pakistan HNMPL - LX3 HEV/ICE 추가',
         'LX3 HEV/ICE 차종 데이터 및 커브 파일 추가'],
        ['2', '17:16', '런타임 에러 방어 코드 추가',
         'Parse→TryParse 81개소, SelectedIndex 체크, null 체크 등 47개 항목'],
        ['3', '20:33', '에러 방어 코드 2차 + CrossThread 수정 + ListBox 누적 방지',
         'CrossThread 방어, 빈 catch 로그, 추가 누락 수정, ListBox 500개 제한'],
    ],
    [1, 2, 8, 5]
)

# ================================================================
# 3. LX3 차종 추가
# ================================================================
doc.add_page_break()
add_heading('3. LX3 차종 추가', 1)

doc.add_paragraph(
    '파키스탄 HNMPL 공장에 LX3 HEV/ICE 2개 차종을 신규 추가하였다. '
    'ECU 통신 클래스 신규 작성, 검사 시퀀스 커브 파일 생성, '
    'DB 모델 데이터 등록, UI 콤보박스 항목 추가를 수행하였다.'
)

add_heading('3.1 ECU 통신 클래스 추가 (cls_ECUs.cs)', 2)
add_table(
    ['항목', 'LX3 HEV', 'LX3 ICE'],
    [
        ['ECU 모델명', 'MOBIS LX3 HEV', 'MOBIS LX3 ICE'],
        ['ECU 타입', 'iMEB2 (현대모비스)', 'MEB5_1 (현대모비스)'],
        ['통신 프로토콜', 'CAN (UDS)', 'CAN (UDS)'],
        ['Send ID', '0x7E7', '0x7D1'],
        ['Receive ID', '0x7EF', '0x7D9'],
        ['SecurityAccess', '불필요 (bypass)', '불필요 (bypass)'],
        ['클래스명', 'MOBIS_LX3H (line 4787~)', 'MOBIS_LX3I (line 5027~)'],
    ],
    [4, 5, 5]
)

add_heading('3.2 구현된 ECU 통신 기능 (14개)', 2)
add_table(
    ['#', '기능', '설명'],
    [
        ['1', 'Start_Communication', 'Extended Diagnostic Session 시작'],
        ['2', 'Stop_Communication', 'Default Session 복귀'],
        ['3', 'ECU_Reset', 'ECU 하드/소프트 리셋'],
        ['4', 'ECU_Identification', 'ECU 버전/파트넘버 읽기'],
        ['5', 'Read__DTC', 'Diagnostic Trouble Code 읽기'],
        ['6', 'Clear_DTC', 'DTC 삭제'],
        ['7', 'Check_Signals', '브레이크 스위치 신호 확인'],
        ['8', 'WSS_Test', '4륜 속도센서 값 읽기'],
        ['9', 'Tester_Present', '세션 유지 (Keep-alive)'],
        ['10', 'Message_Falg', 'ECU 메시지 플래그 확인'],
        ['11', 'Dynamic_Step', 'ABS Dynamic 단계별 제어'],
        ['12', 'Dynamic_Auto', 'ABS Dynamic 자동 시퀀스'],
        ['13', 'ESP_Step', 'ESC/ESP 단계별 제어'],
        ['14', 'ESS_LampTest', 'ESS 경고등 테스트'],
    ],
    [1, 4, 11]
)

add_heading('3.3 ECUs 라우팅 등록 (cls_ECUs.cs)', 2)
doc.add_paragraph(
    'ECUs 클래스의 각 메서드(ECU_Setting, SecurityAccess, Start/Stop_Communication, '
    'ECU_Reset, ECU_Identification, Read/Clear_DTC, Check_Signals, WSS_Test, '
    'Tester_Present, Message_Falg, Dynamic_Step/Auto, ESP_Step, ESS_LampTest)에 '
    'LX3 HEV/ICE case 분기를 추가하여 해당 클래스로 라우팅.'
)

add_heading('3.4 Driving Curve 파일 추가', 2)
add_table(
    ['항목', '내용'],
    [
        ['파일명', 'Curve(LX3).crv'],
        ['위치', 'bin\\Debug\\DCurve\\'],
        ['검사 스텝', '33 Steps'],
        ['총 소요 시간', '167초'],
        ['최대 속도', '100 km/h'],
        ['주요 검사 항목', 'WSS, Speedometer, Cruise, Drag, Brake, ABS Dynamic, Parking'],
    ],
    [4, 12]
)

add_heading('3.5 UI 등록', 2)
add_table(
    ['파일', '내용'],
    [
        ['fomSetup.cs (line 257~258)', 'cbo_ECUs 콤보박스에 MOBIS LX3 HEV/ICE 항목 추가'],
        ['fomDebug.cs (line 231~232)', 'cbo_ECUs 콤보박스에 MOBIS LX3 HEV/ICE 항목 추가'],
        ['fomDebug.cs (line 1541~1552)', 'ESP Step 처리에 LX3 HEV/ICE case 분기 추가'],
        ['DB (KI-RnB.mdb)', 'tbl_CarModel 테이블에 LX3 HEV/ICE 모델 데이터 등록'],
    ],
    [5, 11]
)

# ================================================================
# 4. 신규 기능: Vehicle Balance PLC 출력
# ================================================================
doc.add_page_break()
add_heading('4. 신규 기능: Vehicle Balance PLC 출력', 1)

doc.add_paragraph(
    '차량 모델별 Balance 검사 여부(Y/N)를 DB에서 관리하고, '
    '차량 모델 선택 시 PLC D562[13] 신호로 자동 전송하는 기능을 추가하였다. '
    'W/B(Wheelbase) Select 신호(D562[0~9])와 동일한 레지스터에서 동시 전송된다.'
)

add_heading('4.1 데이터 흐름', 2)
add_table(
    ['단계', '위치', '설명'],
    [
        ['1. 설정', 'fomSetup → chk_balance', '모델별 Balance Y/N을 체크박스로 설정'],
        ['2. DB 저장', 'tbl_CarModel.dbBalance', 'Add/Edit 시 "Y" 또는 "N"으로 저장'],
        ['3. DB 조회', 'clsDBSql → Select/KeyBox/Barcode', '모델 선택 시 dbBalance 값 읽기'],
        ['4. PLC 전송', 'fom_Main → Sel_Vehicles()', 'PLC.DO.Vehicle_Balance 설정 후 PLC_Put_D562() 호출'],
        ['5. PLC 신호', 'D562[13] Vehicle_Balance', 'W/B Select 신호와 동시에 PLC로 전송'],
    ],
    [2, 5, 9]
)

add_heading('4.2 수정 파일', 2)
add_table(
    ['파일', '수정 내용'],
    [
        ['clsDBSql.cs', 'dbBalance 프로퍼티 추가, strModel/str_List 컬럼 추가, Init/조회 4개/Insert/Update 반영'],
        ['fomSetup.cs', 'SelectModel에서 chk_balance 표시, Add_ModelList/EditModelList에서 dbBalance 저장'],
        ['fom_Main.cs', 'Sel_Vehicles에서 PLC.DO.Vehicle_Balance 설정 (PLC_Put_D562와 동시 전송)'],
        ['cls_PLCs.cs', 'Vehicle_Balance 프로퍼티 D562[13], PLC_562_Mapp 매핑 (기 구현)'],
        ['fomSetup.Designer.cs', 'chk_balance CheckBox 컨트롤 (기 구현)'],
    ],
    [4, 12]
)

# ================================================================
# 5. 에러 방어 코드 상세
# ================================================================
doc.add_page_break()
add_heading('5. 에러 방어 코드 상세', 1)

# 3.1
add_heading('5.1 Parse → TryParse 전환 (1차)', 2)
doc.add_paragraph(
    '사용자 입력(TextBox) 및 외부 데이터(시리얼, 파일)를 파싱하는 코드에서 '
    'int.Parse, double.Parse, Convert.ToXxx 호출을 TryParse로 전환하여 '
    'FormatException 방지.'
)
add_table(
    ['구분', '수량', '설명'],
    [
        ['1차 수정 (전체)', '81개소', 'int/double/float Parse → TryParse'],
        ['2차 추가 (fomDebug)', '15개소', 'Convert.ToSingle(txtXXX.Text) → float.TryParse'],
        ['2차 추가 (fomCurve)', '3개소', 'Convert.ToInt32(dgv셀) → int.TryParse'],
    ],
    [4, 3, 9]
)

# 3.2
add_heading('5.2 SelectedIndex / SelectedItem 방어', 2)
add_table(
    ['파일', '수량', '내용'],
    [
        ['fomSetup.cs', '18개소', 'ComboBox.SelectedIndex 대입 시 범위 체크 (>= 0 && < Items.Count)'],
        ['fomDebug.cs', '3개소', 'lst_Step, cboIndex, cboIdent SelectedIndex >= 0 체크'],
        ['fom_Data.cs', '2개소', 'cboModel.SelectedItem null 체크'],
        ['fomCurve.cs', '2개소', 'dgvCurve.CurrentRow null 체크'],
    ],
    [3, 2, 11]
)

# 3.3
add_heading('5.3 CrossThread 방어 (InvokeRequired)', 2)
doc.add_paragraph(
    'SerialPort DataReceived, Thread 등 백그라운드 스레드에서 UI 컨트롤에 '
    '직접 접근하는 코드에 InvokeRequired + BeginInvoke 패턴을 적용.'
)
add_table(
    ['파일', '메서드', '변경 내용'],
    [
        ['fom_Main.cs', 'Indi_LogData', 'InvokeRequired + BeginInvoke. chk_Indi.Checked 직접 접근 제거'],
        ['fom_Main.cs', 'ABSB_LogData', '동일 패턴. chk_ABSB.Checked 직접 접근 제거'],
        ['fom_Main.cs', 'PLC_Log_Data', '동일 패턴. chk_PLCs.Checked 직접 접근 제거'],
        ['fom_Main.cs', 'ABSBScanHerz', 'InvokeRequired + BeginInvoke. catch 블록 UI 접근 제거'],
        ['fom_Test.cs', 'Refresh_Data 외 6개', '7개 public 메서드에 InvokeRequired 방어 추가'],
        ['cls_ABSB.cs', 'ThreadABSB 외', 'lbl_Ctrl.ForeColor 접근 4곳 BeginInvoke + IsHandleCreated'],
    ],
    [3, 4, 9]
)

# 3.4
add_heading('5.4 NullReference 방어', 2)
add_table(
    ['파일', '메서드', '내용'],
    [
        ['clsDBSql.cs', 'DS_Select', 'DBs_Conn == null 시 빈 DataSet 반환'],
        ['clsDBSql.cs', 'DT_Select', 'DBs_Conn == null 시 빈 DataTable 반환'],
        ['clsDBSql.cs', 'Execute', 'DBs_Conn == null 시 return'],
        ['clsDBSql.cs', 'Create', 'DBs_Conn == null 시 return'],
    ],
    [3, 3, 10]
)

# 3.5
add_heading('5.5 빈 catch 블록 로그 추가', 2)
doc.add_paragraph(
    '프로젝트 전체 170개 catch 블록 중 실제로 비어있는 36개에 '
    'Logs.MakeLog_File(Log_His.Err_, "메서드명: " + ex.Message) 로그를 추가. '
    'cls_Logs.cs의 ExceptionErr (재귀 방지용) 1개는 의도적으로 제외.'
)

add_table(
    ['분류', '수량', '대표 파일'],
    [
        ['장비 통신 (시리얼/NI/ECU)', '9개', 'clsBS205, clsDAQmx, clsNeoVI, cls_ABSB, cls_H2Ys'],
        ['파일/DB I/O', '3개', 'clsDBSql, cls_PSet'],
        ['UI 이벤트 (폼/컨트롤)', '10개', 'fomSetup, fomDebug, fom_Keys, fom1Gage, fom4Gage'],
        ['CrossThread Invoke 방어', '8개', 'fom_Main (Prog_LogData, PLC, ABSB, Indi, DLC, TestImg)'],
        ['검사 실행', '2개', 'cls_Test (Scan_Sensors, Ret_ECU_Logs)'],
        ['기타', '4개', 'fom_Load, fom_Loss, fom_Test, Chery_TestPresentThread'],
    ],
    [5, 2, 9]
)

# 3.6
add_heading('5.6 ListBox 누적 방지', 2)
doc.add_paragraph(
    '검사 반복 시 ListBox 항목이 무한 누적되어 UI 성능이 저하되고, '
    'Invoke 데드락으로 바코드 수신이 중단되는 문제를 방지하기 위해 '
    '500개 초과 시 오래된 항목을 자동 삭제하도록 수정.'
)
add_table(
    ['ListBox', '메서드', '용도'],
    [
        ['lst_Logs', 'Prog_LogData', '프로그램 로그'],
        ['lst_PLCs', 'PLC_Log_Data', 'PLC 통신 로그'],
        ['lst_ABSB', 'ABSB_LogData', 'ABS 보드 로그'],
        ['lst_Indi', 'Indi_LogData', '인디케이터 로그'],
        ['lst_DLCs', 'DLCs_LogData', 'ECU 통신 로그'],
    ],
    [3, 4, 9]
)
doc.add_paragraph('※ 삭제된 로그는 파일(Log File.log)에 별도 기록되므로 데이터 유실 없음.')

# ================================================================
# 4. 기타 수정
# ================================================================
doc.add_page_break()
add_heading('6. 기타 수정', 1)

add_heading('6.1 바코드 VIN 우선 적용', 2)
doc.add_paragraph(
    'txtVinNo를 Key_Vehicles 호출 전에 설정하여 H2Y.Sleep 중 DoEvents 경쟁 방지. '
    '바코드 스캔 시 새 VIN이 항상 우선 적용.'
)
add_bullet('원인: ', 'Key_Vehicles → Sel_Vehicles → H2Y.Sleep(DoEvents) 중 타이머 틱이 끼어들어 이전 VIN으로 검사 시작')
add_bullet('수정: ', 'txtVinNo.Text = pVinNo를 Key_Vehicles 호출 전에 배치')

add_heading('6.2 WSS 파싱 오프셋 수정', 2)
doc.add_paragraph(
    'LX3 HEV(iMEB2)와 ICE(MEB5_1)의 WSS DID(01 04) 응답 길이가 다름. '
    'HEV: 0x2E(46바이트) → Ident[14~17], ICE: 0x2C(44바이트) → Ident[15~18].'
)

add_heading('6.3 NRC 0x78 pending 처리', 2)
doc.add_paragraph(
    'clsNeoVI.cs에서 ECU 응답 7F XX 78 (requestCorrectlyReceivedResponsePending)을 '
    '실패 대신 다음 응답 대기로 처리. LX3 DTC Clear 시 X 판정되던 문제 해결. '
    'NeoVI 공통 코드이므로 모든 ECU에 적용.'
)

add_heading('6.4 PLC Cancel/Stop 시 VIN 숨김', 2)
doc.add_paragraph(
    '검사 시작 전 PLC Cancel/Stop 시 FomFlash.VinNo_Hide() 호출. '
    'OrderStopped()와 타이머 틱에 추가.'
)

add_heading('6.5 dgv_List 더블클릭 에러 및 재진입 방지', 2)
doc.add_paragraph(
    'CellDoubleClick과 CurrentCellChanged 이벤트가 동시에 SelectResult를 호출하여 '
    'SetCurrentCellAddressCore 재진입 에러 발생. CellDoubleClick 이벤트 등록 제거 및 '
    'SelectResult를 BeginInvoke로 감싸서 DataGridView 셀 변경 완료 후 실행되도록 수정.'
)

add_heading('6.6 Cancel/Stop 행 표시 (dbStopFlag)', 2)
doc.add_paragraph(
    'tbl_InfoData에 dbStopFlag 컬럼 추가. 검사 종료 시 TSet.StopFlag 값을 DB에 저장. '
    'dgv_List에서 StopFlag > 0인 행의 VINNO 텍스트를 빨간색으로 표시. '
    '첫 번째 Info_DataAdd 호출 전 TSet.StopFlag = 0 초기화 추가.'
)

add_heading('6.8 DB 컬럼명 수정 및 기타', 2)
add_bullet('', 'dbBalance → dbCarBalance: DB 컬럼명과 코드 일치시킴')
add_bullet('', 'dgvModel CurrentRow null 체크 추가 (fomSetup dgvModel_CurrentCellChanged)')

add_heading('6.9 cboModel 드롭다운 닫힘 방지', 2)
doc.add_paragraph(
    '현장 PC에서 cboModel 드롭다운을 열면 목록이 나타났다가 바로 사라지는 현상 대응. '
    'tmr_Main 타이머(300ms)에서 PLC Select 신호 처리 시 cboModel.Text를 변경하면 '
    '열려있는 드롭다운이 강제로 닫히는 문제.'
)
add_bullet('원인: ', 'PLC Select 신호 → Key_Vehicles() → cboModel.Text 변경 → 드롭다운 닫힘')
add_bullet('수정: ', '!cboModel.DroppedDown 조건 추가. 드롭다운이 열린 동안 PLC Select 처리 지연')
add_bullet('영향: ', '드롭다운 닫힌 후 다음 틱(0.3초 이내)에서 정상 처리. 패널티 없음')

add_heading('6.10 fomCurve 새 커브 타이틀 표시', 2)
doc.add_paragraph(
    '새 Driving Curve 생성 시 fomCurve 타이틀바에 모델명이 표시되지 않던 문제 수정. '
    'crv_Mode=1 (새로 만들기) 분기에 this.Text 설정 추가.'
)

add_heading('6.11 CAN 멀티프레임 Return 설정', 2)
doc.add_paragraph(
    'STD_CAN_Read() First Frame 핸들러에서 긍정응답(GetD[2] != "7F") 시 '
    'Return = true 설정 누락 수정. HEV DTC Read가 멀티프레임으로 수신되면서 '
    'Return이 미설정되어 X 판정되던 문제 해결. ICE는 싱글프레임이라 영향 없었음.'
)

add_heading('6.12 바코드 자릿수 검증', 2)
doc.add_paragraph(
    '바코드 스캔 시 자릿수 검증 및 DB 매칭 실패 시 팝업 표시 추가. '
    '9자리 미만: too short 팝업, 17자리 초과: too long 팝업, '
    'DB 매칭 실패(Count=0): 스캔값 표시 및 Setting 확인 안내.'
)

add_heading('6.13 icsNeoClass catch 타입 복원', 2)
doc.add_paragraph(
    'ConvertFromHex 메서드의 catch가 OverflowException에서 Exception으로 '
    '확대되었던 것을 System.OverflowException으로 복원. '
    'FormatException 등 다른 예외가 삼켜지는 사이드 이펙트 방지.'
)


# ================================================================
# 5. 사이드 이펙트 검증 결과
# ================================================================
add_heading('7. 사이드 이펙트 검증 결과', 1)

add_table(
    ['검증 항목', '결과', '비고'],
    [
        ['Invoke → BeginInvoke (4개 메서드)', '안전', 'fire-and-forget 로깅. 호출 후 UI 의존 없음'],
        ['fom_Test InvokeRequired (7개 메서드)', '안전', '현재 UI 스레드에서만 호출됨'],
        ['cls_ABSB BeginInvoke + IsHandleCreated', '안전', '4곳 모두 보호 확인'],
        ['clsDBSql null 체크', '안전', '빈 DataSet/DataTable 반환. 호출자 처리 가능'],
        ['icsNeoClass catch 타입', '복원 완료', 'OverflowException으로 원상복구'],
        ['fomCurve TryParse', '낮은 위험', 'DataGridView 셀 비정상 확률 극히 낮음'],
        ['fomSetup SelectedIndex 범위 체크', '안전', '범위 밖이면 스킵 (기본값 유지)'],
        ['빈 catch 로그 추가 (36개)', '안전', '기존 흐름 변경 없음. 로그만 추가'],
        ['ListBox 500개 제한', '안전', '삭제 로그는 파일에 별도 기록'],
    ],
    [6, 3, 7]
)

# ================================================================
# 6. 수정 파일 목록
# ================================================================
add_heading('8. 수정 파일 목록 (25개)', 1)

add_table(
    ['#', '파일명', '수정 내용 요약'],
    [
        ['1', 'Chery_TestPresentThread.cs', '빈 catch 로그 추가'],
        ['2', 'clsBS205.cs', '빈 catch 로그 추가'],
        ['3', 'clsDAQmx.cs', '빈 catch 로그 추가 (3곳)'],
        ['4', 'clsDBSql.cs', 'DBs_Conn null 체크 + 빈 catch 로그'],
        ['5', 'clsNeoVI.cs', '빈 catch 로그 추가 (2곳)'],
        ['6', 'cls_ABSB.cs', 'CrossThread BeginInvoke + IsHandleCreated + 빈 catch 로그'],
        ['7', 'cls_H2Ys.cs', '빈 catch 로그 추가'],
        ['8', 'cls_PSet.cs', '빈 catch 로그 추가 (2곳)'],
        ['9', 'cls_Test.cs', '빈 catch 로그 추가 (2곳)'],
        ['10', 'fom1Gage.cs', '빈 catch 로그 추가'],
        ['11', 'fom4Gage.cs', '빈 catch 로그 추가'],
        ['12', 'fomCurve.cs', 'TryParse + CurrentRow null 체크 + 새 커브 타이틀 표시'],
        ['13', 'fomDebug.cs', 'TryParse 15개 + SelectedIndex 체크 3개 + 빈 catch 로그'],
        ['14', 'fomSetup.cs', 'SelectedIndex 범위 체크 18개 + 빈 catch 로그 3개'],
        ['15', 'fom_Data.cs', 'SelectedItem null 체크 2개'],
        ['16', 'fom_Keys.cs', '빈 catch 로그 추가'],
        ['17', 'fom_Load.cs', '빈 catch 로그 추가'],
        ['18', 'fom_Loss.cs', '빈 catch 로그 추가'],
        ['19', 'fom_Main.cs', 'CrossThread 4개 + ListBox 제한 5개 + DroppedDown + 배포 원복'],
        ['20', 'fom_Test.cs', 'InvokeRequired 7개 + 빈 catch 로그'],
        ['21', 'icsNeoClass.cs', 'catch 타입 복원 + 빈 catch 로그'],
        ['22', 'MachineSet.def', 'OnwerDrv=0 → 1 (배포용)'],
        ['23', 'LX3_ABS_Test_Sequence.pptx', 'LX3 ABS 검사 시퀀스 PPT (신규)'],
        ['24', 'make_ppt.py', 'PPT 생성 스크립트 (신규)'],
    ],
    [1, 5, 10]
)

output = r"E:\파키스탄 LX3\KI-RnB_1208\KI-RnB_작업보고서_20260815.docx"
doc.save(output)
print(f"Done: {output}")
